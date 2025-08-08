import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

CONTENT = [
    ("Knowledge Bot Widget Integration Guide", "h1"),
    ("Overview", "h2"),
    ("Embed a glassmorphism chat widget on any website. It preserves token handling, session management, markdown rendering, DOM sanitization, and streaming.", "p"),
    ("Prerequisites", "h2"),
    ("- Token from CMS → Bot Management → copy the bot’s widget token\n- Optional: DOMPurify for sanitization\n- Recommended meta for mobile: <meta name=\"viewport\" content=\"width=device-width, initial-scale=1, viewport-fit=cover\">", "pre"),
    ("Production embed", "h2"),
    ("""
<meta name=\"viewport\" content=\"width=device-width, initial-scale=1, viewport-fit=cover\" />
<script src=\"https://cdn.jsdelivr.net/npm/dompurify@3.0.2/dist/purify.min.js\" defer></script>
<script
  src=\"https://knowledge-bot-retrieval.onrender.com/static/widget.js\"
  data-token=\"YOUR_WIDGET_TOKEN_HERE\"
  data-bubble-title=\"Ask ClearlyClear\"
  data-search-title=\"ClearlyClear Concierge\"
  data-show-logo=\"false\"
  defer
></script>
""".strip(), "code"),
    ("Local development", "h2"),
    ("""
<script
  src=\"http://localhost:8000/static/widget.js\"
  data-token=\"YOUR_LOCAL_TEST_TOKEN\"
  data-bubble-title=\"Chat\"
  data-search-title=\"Assistant\"
  data-show-logo=\"false\"
  defer
></script>
""".strip(), "code"),
    ("Notes", "h3"),
    ("- Auto-detects localhost/127.0.0.1 for API calls; otherwise production.\n- Adjust the script URL if you use a different local port.", "p"),
    ("Attributes", "h2"),
    ("- data-token (required): your bot’s widget token\n- data-bubble-title (optional): text in bubble; if omitted, an icon is shown\n- data-search-title (optional): header title; blank if omitted\n- data-show-logo (optional): \"true\" to show logo; hidden by default", "p"),
    ("Behavior", "h2"),
    ("- Bottom-right bubble; panel grows upward only; capped to screen height (vh/dvh)\n- Enter to send; or click Send\n- Streams responses via SSE/JSON; control frames filtered\n- Markdown rendered; sources in a collapsible ‘Show Sources’ section\n- If DOMPurify present, all rendered content sanitized", "p"),
    ("Styling and isolation", "h2"),
    ("- Scoped .kb-* styles\n- Ensure global resets don’t override widget inputs/buttons", "p"),
    ("Mobile specifics", "h2"),
    ("- viewport-fit=cover and dvh when supported\n- Input font-size 16px on small screens to avoid iOS zoom\n- Header/close always accessible; content area scrolls", "p"),
    ("Token source", "h2"),
    ("- CMS → Bot Management → select bot → copy token → set data-token\n- Rotating tokens: update the attribute only", "p"),
    ("Troubleshooting", "h2"),
    ("- White-on-white text: ensure site CSS doesn’t force colors on .kb-*\n- CORS: allow your site’s origin on the backend\n- Local connection: server must run at http://localhost:8000", "p"),
    ("Security", "h2"),
    ("- Multi-tenant isolation enforced server-side\n- Client sanitization via DOMPurify", "p"),
]

def add_heading(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level)

def add_paragraph(doc: Document, text: str, mono: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if mono:
        run.font.name = 'Courier New'
    run.font.size = Pt(11)


def main():
    doc = Document()
    # Title
    doc.add_heading(CONTENT[0][0], 0)

    for text, kind in CONTENT[1:]:
        if kind == 'h2':
            add_heading(doc, text, 1)
        elif kind == 'h3':
            add_heading(doc, text, 2)
        elif kind in ('code', 'pre'):
            add_paragraph(doc, text, mono=True)
        else:
            add_paragraph(doc, text)

    out_path = os.path.join(os.getcwd(), 'Knowledge Bot Widget Integration Guide (Updated).docx')
    doc.save(out_path)
    print(f'Wrote: {out_path}')

if __name__ == '__main__':
    main() 